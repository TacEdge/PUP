#!/usr/bin/env python3
"""
Build the NZ Army 'Performance Under Pressure and Combat Mindset' capability
note as a branded .docx staff paper.

Layout/typesetting only — the source text in SOURCE_FILE is reproduced
verbatim. The single permitted editorial change: inline filename citations
(e.g. 'PUP Workbook.pdf') are lifted out of the running text and replaced
with sequential superscript markers, collected in a 'Source Documents'
list at the end.
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./capability-note.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/capability-note.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"
DOCUMENT_REFERENCE = "NZALC 2026"
DATE               = "July 2026"
ORIGINATOR         = "Major M Coom, Chief Instructor, New Zealand Army Leadership Centre"
VERSION            = "Draft v0.1"
DISTRIBUTION       = "ACS; Comd ATG; Comd NZALC"
LIST_STYLE         = "bullets"

TITLE    = "Performance Under Pressure and Combat Mindset in the New Zealand Army"
SUBTITLE = "Draft Capability Note"
FOOTER_LEFT = "PUP & Combat Mindset — Capability Note"

# Brand palette (NZDF Visual Identity Standards — NZ Army)
ARMY_RED       = "C62026"
DARKEST_HOUR   = "000000"
RUAPEHU_WHITE  = "FFFFFF"
SWAMP_GREEN    = "002516"
KAWAKAWA_LEAF  = "3A4B00"
WAIOURU_HILLS  = "A89662"
MOAWHANGO      = "CDD2B7"

FONT_HEAD = "Arial"         # brand face: Haas Grotesk
FONT_BODY = "Book Antiqua"  # brand face: Publico (fallback Georgia)
FONT_MONO = "Consolas"

# Known source-document filenames to convert to reference markers.
CITATION_FILES = [
    "PUP Workbook.pdf",
    "LSYS (Officers) Workbook_May 26.pdf",
    "ELDA Lead Leaders Workbook (LPR).pdf",
    "Lead Self Workbook (TAD) 2026.pdf",
    "Lead Teams Workbook.pdf",
    "LDS LEAD LEADERS WORKBOOK.pdf",
]
CITATION_RE = re.compile(
    r"\s*(" + "|".join(re.escape(f) for f in CITATION_FILES) + r")"
)

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _pPr(p).append(shd)


def set_border(p, side, color, sz, space=4):
    """Paragraph border. sz is in eighths of a point."""
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    """Set a font name and strip any theme-font attributes that would
    otherwise take precedence in Word."""
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    """Insert a Word field (PAGE, NUMPAGES, TOC ...) into a paragraph."""
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ------------------------------------------------------- citation handling --

footnote_order = []  # unique filenames in order of first appearance


def citation_number(fname):
    if fname not in footnote_order:
        footnote_order.append(fname)
    return footnote_order.index(fname) + 1


def split_citations(text):
    """Return list of (kind, value) segments: ('text', str) / ('marker', n)."""
    segments = []
    pos = 0
    for m in CITATION_RE.finditer(text):
        if m.start() > pos:
            segments.append(("text", text[pos:m.start()]))
        segments.append(("marker", citation_number(m.group(1))))
        pos = m.end()
    if pos < len(text):
        segments.append(("text", text[pos:]))
    return segments


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    """Add runs for text, honouring **bold** spans and citation markers."""
    for kind, value in split_citations(text):
        if kind == "marker":
            run = p.add_run(str(value))
            run.font.superscript = True
            force_font(run, base_font)
            if size:
                run.font.size = size
            continue
        # split on **bold** spans
        idx = 0
        for m in BOLD_RE.finditer(value):
            if m.start() > idx:
                _run(p, value[idx:m.start()], base_font, size, bold, italic, color)
            _run(p, m.group(1), base_font, size, True, italic, color)
            idx = m.end()
        if idx < len(value):
            _run(p, value[idx:], base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

# Page geometry — A4
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(2.4)
section.bottom_margin = Cm(2.4)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.header_distance = Cm(1.1)
section.footer_distance = Cm(1.1)
section.different_first_page_header_footer = True

TEXT_WIDTH_CM = 16.0  # 21 - 2.5 - 2.5

# Base style — nothing may fall back to Calibri
normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(11)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.15
nf.space_after = Pt(6)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Section headings (Heading 1 keeps outline level -> TOC)
h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_HEAD)
h1.font.bold = True
h1.font.size = Pt(14)
force_color(h1, SWAMP_GREEN)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.keep_with_next = True

# Sub-headings
h2 = doc.styles["Heading 2"]
strip_style_rpr(h2)
force_font(h2, FONT_HEAD)
h2.font.bold = True
h2.font.size = Pt(11)
force_color(h2, DARKEST_HOUR)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.keep_with_next = True


def add_section_heading(text):
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "left", ARMY_RED, 24, space=8)  # Army Red section tab
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(14), bold=True,
                  color=SWAMP_GREEN)
    return p


def add_sub_heading(text):
    p = doc.add_paragraph(style="Heading 2")
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(11), bold=True,
                  color=DARKEST_HOUR)
    return p


def add_body(text):
    p = doc.add_paragraph()
    add_text_runs(p, text)
    return p


def add_display(text_lines):
    """Single-line display statement(s): indented, italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(text_lines):
        if i:
            p.add_run().add_break()
        add_text_runs(p, line, italic=True)
    return p


def add_callout(text_lines):
    """Key summation callout: Moawhango fill, Army Red left border, bold."""
    p = doc.add_paragraph()
    set_shading(p, MOAWHANGO)
    set_border(p, "left", ARMY_RED, 28, space=8)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    for i, line in enumerate(text_lines):
        if i:
            p.add_run().add_break()
        font = FONT_HEAD if "→" in line else FONT_BODY
        add_text_runs(p, line, base_font=font, bold=True)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    bullet = p.add_run("•\t")
    force_font(bullet, FONT_HEAD)
    force_color(bullet, ARMY_RED)
    bullet.bold = True
    add_text_runs(p, text)
    return p


def add_numbered(num, lead, continuation):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    n = p.add_run(f"{num}.\t")
    force_font(n, FONT_HEAD)
    n.bold = True
    add_text_runs(p, lead)
    if continuation:
        c = doc.add_paragraph()
        c.paragraph_format.left_indent = Cm(0.75)
        c.paragraph_format.space_after = Pt(6)
        add_text_runs(c, continuation)
    return p


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    set_border(p, "bottom", WAIOURU_HILLS, 6, space=1)
    return p


def add_page_break():
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


# Cover (first page): marking top and bottom only
marking_paragraph(section.first_page_header)
marking_paragraph(section.first_page_footer)

# Content pages: marking top centre
marking_paragraph(section.header)

# Content pages footer: marking centre + info line
marking_paragraph(section.footer)
info = section.footer.add_paragraph()
info.paragraph_format.space_before = Pt(2)
info.paragraph_format.space_after = Pt(0)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)


def footer_run(text, bold=False):
    run = info.add_run(text)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(8.5)
    run.bold = bold
    force_color(run, DARKEST_HOUR)
    return run


footer_run(FOOTER_LEFT)
footer_run("\t")
footer_run(DOCUMENT_REFERENCE)
footer_run("\tPage ")
for r in add_field(info, "PAGE", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)
footer_run(" of ")
for r in add_field(info, "NUMPAGES", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)

# ------------------------------------------------------------------ cover ---

doc.add_picture(LOGO_FILE, width=Mm(60))

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(96)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(10)
set_border(title_p, "bottom", ARMY_RED, 20, space=10)
t = title_p.add_run(TITLE)
force_font(t, FONT_HEAD)
t.font.size = Pt(26)
t.bold = True
force_color(t, DARKEST_HOUR)

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(6)
sub_p.paragraph_format.space_after = Pt(120)
s = sub_p.add_run(SUBTITLE)
force_font(s, FONT_HEAD)
s.font.size = Pt(13)
force_color(s, SWAMP_GREEN)

meta_rows = [
    ("Reference", DOCUMENT_REFERENCE),
    ("Date", DATE),
    ("Originator", ORIGINATOR),
    ("Version", VERSION),
    ("Distribution", DISTRIBUTION),
]
for i, (label, value) in enumerate(meta_rows):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(3.2))
    if i == 0:
        set_border(p, "top", WAIOURU_HILLS, 6, space=6)
    if i == len(meta_rows) - 1:
        set_border(p, "bottom", WAIOURU_HILLS, 6, space=6)
    lab = p.add_run(f"{label}\t")
    force_font(lab, FONT_HEAD)
    lab.font.size = Pt(10)
    lab.bold = True
    force_color(lab, SWAMP_GREEN)
    val = p.add_run(value)
    force_font(val, FONT_HEAD)
    val.font.size = Pt(10)
    force_color(val, DARKEST_HOUR)

add_page_break()

# --------------------------------------------------------------- contents ---

toc_head = doc.add_paragraph()
toc_head.paragraph_format.space_after = Pt(12)
set_border(toc_head, "left", ARMY_RED, 24, space=8)
r = toc_head.add_run("Contents")
force_font(r, FONT_HEAD)
r.font.size = Pt(14)
r.bold = True
force_color(r, SWAMP_GREEN)

toc_p = doc.add_paragraph()
add_field(toc_p, r'TOC \o "1-1" \h \z \u',
          "Update fields to generate the table of contents.")

add_page_break()

# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    if not line.strip():
        i += 1
        continue

    if line.startswith("# ") or line.startswith("*Draft Capability Note*"):
        i += 1  # title/subtitle already on cover
        continue

    if line.strip() == "---":
        add_divider()
        i += 1
        continue

    if line.startswith("## "):
        add_section_heading(line[3:].strip())
        i += 1
        continue

    if line.startswith("### "):
        add_sub_heading(line[4:].strip())
        i += 1
        continue

    if line.startswith(">! "):
        block = []
        while i < len(lines) and lines[i].startswith(">! "):
            block.append(lines[i][3:].strip())
            i += 1
        add_callout(block)
        continue

    if line.startswith("> "):
        block = []
        while i < len(lines) and lines[i].startswith("> "):
            block.append(lines[i][2:].strip())
            i += 1
        add_display(block)
        continue

    if line.startswith("* "):
        add_bullet(line[2:].strip())
        i += 1
        continue

    m = re.match(r"^(\d+)\. (.*)$", line)
    if m:
        num, lead = m.group(1), m.group(2).strip()
        continuation = None
        if i + 1 < len(lines) and lines[i + 1].startswith("   ") \
                and lines[i + 1].strip():
            continuation = lines[i + 1].strip()
            i += 1
        add_numbered(num, lead, continuation)
        i += 1
        continue

    add_body(line.strip())
    i += 1

# -------------------------------------------------------- source documents --

add_divider()
add_section_heading("Source Documents")

for n, fname in enumerate(footnote_order, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    numr = p.add_run(f"{n}.\t")
    force_font(numr, FONT_HEAD)
    numr.bold = True
    fn = p.add_run(fname)
    force_font(fn, FONT_MONO)
    fn.font.size = Pt(10)
    tail = p.add_run(" — internal NZALC workbook.")
    force_font(tail, FONT_BODY)

# ----------------------------------------------------------------- finish ---

# Prune unused template styles so no off-palette colour or non-brand font
# remains anywhere in the file, and point document defaults at the body font.
USED_STYLE_IDS = {"Normal", "Heading1", "Heading2", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

mark_update_fields(doc)

props = doc.core_properties
props.title = TITLE
props.author = ORIGINATOR
props.subject = SUBTITLE

doc.save(OUTPUT_DOCX)

print(f"Saved {OUTPUT_DOCX}")
print(f"Footnote markers reference {len(footnote_order)} unique source documents:")
for n, fname in enumerate(footnote_order, start=1):
    print(f"  {n}. {fname}")
