#!/usr/bin/env python3
"""
Build the NZ Army 'Adventure Race — Participant Withdrawal SOP' as a compact
branded .docx. Letterhead-style first page (no cover, no TOC); body typeset
from SOURCE_FILE on the established brand system.
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./withdrawal-sop.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/withdrawal-sop.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"
DOCUMENT_REFERENCE = "ELDA Lead Systems Course Workbook, pp 60–62 (Learning Objectives)"
FOOTER_REFERENCE   = "ACS 2026"
DATE               = "July 2026"
ORIGINATOR         = "NZ Army Leadership Centre | Army Command School"
VERSION            = "Draft v0.1"

TITLE       = "Adventure Race — Participant Withdrawal SOP"
FOOTER_LEFT = "Participant Withdrawal SOP"

ARMY_RED      = "C62026"
DARKEST_HOUR  = "000000"
RUAPEHU_WHITE = "FFFFFF"
SWAMP_GREEN   = "002516"
WAIOURU_HILLS = "A89662"
MOAWHANGO     = "CDD2B7"

FONT_HEAD = "Arial"
FONT_BODY = "Book Antiqua"

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _pPr(p).append(shd)


def set_border(p, side, color, sz, space=4):
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    idx = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > idx:
            _run(p, text[idx:m.start()], base_font, size, bold, italic, color)
        _run(p, m.group(1), base_font, size, True, italic, color)
        idx = m.end()
    if idx < len(text):
        _run(p, text[idx:], base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.3)
section.right_margin = Cm(2.3)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

TEXT_WIDTH_CM = 16.4

normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(10.5)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.12
nf.space_after = Pt(5)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_HEAD)
h1.font.bold = True
h1.font.size = Pt(12)
force_color(h1, SWAMP_GREEN)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True


def add_section_heading(text):
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "left", ARMY_RED, 24, space=8)
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(12), bold=True,
                  color=SWAMP_GREEN)
    return p


def add_body(text):
    p = doc.add_paragraph()
    add_text_runs(p, text)
    return p


def add_callout(text_lines):
    p = doc.add_paragraph()
    set_shading(p, MOAWHANGO)
    set_border(p, "left", ARMY_RED, 28, space=8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(text_lines):
        if i:
            p.add_run().add_break()
        add_text_runs(p, line, bold=True)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.42)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.7))
    bullet = p.add_run("•\t")
    force_font(bullet, FONT_HEAD)
    force_color(bullet, ARMY_RED)
    bullet.bold = True
    add_text_runs(p, text)
    return p


def add_numbered(num, lead):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.7))
    n = p.add_run(f"{num}.\t")
    force_font(n, FONT_HEAD)
    n.bold = False
    add_text_runs(p, lead)
    return p


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


marking_paragraph(section.header)
marking_paragraph(section.footer)
info = section.footer.add_paragraph()
info.paragraph_format.space_before = Pt(2)
info.paragraph_format.space_after = Pt(0)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)


def footer_run(text):
    run = info.add_run(text)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(8.5)
    force_color(run, DARKEST_HOUR)
    return run


footer_run(FOOTER_LEFT)
footer_run("\t")
footer_run(FOOTER_REFERENCE)
footer_run("\tPage ")
for r in add_field(info, "PAGE", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)
footer_run(" of ")
for r in add_field(info, "NUMPAGES", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)

# -------------------------------------------------------------- letterhead --

_logo = Image.open(LOGO_FILE).convert("RGBA")
_logo = _logo.crop(_logo.getchannel("A").getbbox())
_buf = io.BytesIO()
_logo.save(_buf, "PNG")
_buf.seek(0)
doc.add_picture(_buf, width=Mm(42))
logo_para = doc.paragraphs[-1]
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after = Pt(14)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(2)
set_border(title_p, "bottom", ARMY_RED, 18, space=8)
t = title_p.add_run(TITLE)
force_font(t, FONT_HEAD)
t.font.size = Pt(19)
t.bold = True
force_color(t, DARKEST_HOUR)

org_p = doc.add_paragraph()
org_p.paragraph_format.space_before = Pt(6)
org_p.paragraph_format.space_after = Pt(6)
o = org_p.add_run(ORIGINATOR)
force_font(o, FONT_HEAD)
o.font.size = Pt(10.5)
force_color(o, SWAMP_GREEN)

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_after = Pt(10)
set_border(meta_p, "bottom", WAIOURU_HILLS, 6, space=6)
for i, (label, value) in enumerate([
        ("Reference", DOCUMENT_REFERENCE), ("Date", DATE),
        ("Version", VERSION)]):
    if i:
        gap = meta_p.add_run("      ")
        force_font(gap, FONT_HEAD)
    lab = meta_p.add_run(f"{label}  ")
    force_font(lab, FONT_HEAD)
    lab.font.size = Pt(9)
    lab.bold = True
    force_color(lab, SWAMP_GREEN)
    val = meta_p.add_run(value)
    force_font(val, FONT_HEAD)
    val.font.size = Pt(9)
    force_color(val, DARKEST_HOUR)

# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    if not line.strip():
        i += 1
        continue

    if line.startswith("# "):
        i += 1  # document title already in letterhead
        continue

    if line.startswith("## "):
        add_section_heading(line[3:].strip())
        i += 1
        continue

    if line.startswith(">! "):
        block = []
        while i < len(lines) and lines[i].startswith(">! "):
            block.append(lines[i][3:].strip())
            i += 1
        add_callout(block)
        continue

    if line.startswith("* "):
        add_bullet(line[2:].strip())
        i += 1
        continue

    m = re.match(r"^(\d+)\. (.*)$", line)
    if m:
        add_numbered(m.group(1), m.group(2).strip())
        i += 1
        continue

    add_body(line.strip())
    i += 1

# ----------------------------------------------------------------- finish ---

USED_STYLE_IDS = {"Normal", "Heading1", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

while doc.paragraphs and not doc.paragraphs[-1].text.strip() \
        and not doc.paragraphs[-1]._p.findall(qn("w:r") + "/" + qn("w:drawing")):
    _el = doc.paragraphs[-1]._element
    _el.getparent().remove(_el)

mark_update_fields(doc)

props = doc.core_properties
props.title = TITLE
props.author = ORIGINATOR

doc.save(OUTPUT_DOCX)
print(f"Saved {OUTPUT_DOCX}")
