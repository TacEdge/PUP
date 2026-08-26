#!/usr/bin/env python3
"""
Build the 'ALC Health & Safety Inspection — Action List' as a compact branded
.docx on the established brand system. Single-page action document: numbered,
closeable actions only; no assessment narrative.
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./hs-inspection-actions.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/hs-inspection-actions.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"
FOOTER_REFERENCE   = "ACS 2026"
DATE               = "August 2026"
ORIGINATOR         = "NZ Army Leadership Centre | Army Command School"

TITLE         = "ALC Health & Safety Inspection"
SUBTITLE_LINE = "Action List"
FOOTER_LEFT   = "ALC Health & Safety Inspection — Action List"

ARMY_RED      = "C62026"
DARKEST_HOUR  = "000000"
SWAMP_GREEN   = "002516"
WAIOURU_HILLS = "A89662"
MOAWHANGO     = "CDD2B7"

FONT_HEAD    = "Aptos"
FONT_BODY    = "Aptos"
FONT_DISPLAY = "Aptos Display"      # title only
FONT_SYMBOL  = "Segoe UI Symbol"    # carries the U+2610 tick box glyph

TICK_BOX = "☐"                 # ☐ BALLOT BOX

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_border(p, side, color, sz, space=4):
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    idx = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > idx:
            _run(p, text[idx:m.start()], base_font, size, bold, italic, color)
        _run(p, m.group(1), base_font, size, True, italic, color)
        idx = m.end()
    if idx < len(text):
        _run(p, text[idx:], base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.3)
section.right_margin = Cm(2.3)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

TEXT_WIDTH_CM = 16.4

normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(10.5)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.12
nf.space_after = Pt(5)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_HEAD)
h1.font.bold = True
h1.font.size = Pt(12)
force_color(h1, SWAMP_GREEN)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True


def add_section_heading(text):
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "left", ARMY_RED, 24, space=8)
    add_text_runs(p, text.upper(), base_font=FONT_HEAD, size=Pt(12), bold=True,
                  color=SWAMP_GREEN)
    return p


def add_body(text):
    p = doc.add_paragraph()
    add_text_runs(p, text)
    return p


def add_action(text):
    """A closeable action, led by a tick box to mark off on the walk-round."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.85)
    p.paragraph_format.first_line_indent = Cm(-0.85)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.85))
    box = p.add_run(TICK_BOX)
    force_font(box, FONT_SYMBOL)
    box.font.size = Pt(14)
    force_color(box, SWAMP_GREEN)
    tab = p.add_run("\t")
    force_font(tab, FONT_BODY)
    add_text_runs(p, text)
    return p


def add_signoff_strip(labels):
    """Sign-off fields side by side: label above an open rule, one per column.

    Laid out across a single strip rather than stacked so the close-out block
    stays on the same page as the actions it closes.
    """
    table = doc.add_table(rows=2, cols=len(labels))
    table.autofit = False
    col_w = Cm(TEXT_WIDTH_CM / len(labels))

    for i, label in enumerate(labels):
        lab_cell = table.cell(0, i)
        lab_cell.width = col_w
        p = lab_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        force_font(run, FONT_HEAD)
        run.font.size = Pt(9)
        run.bold = True
        force_color(run, SWAMP_GREEN)

        rule_cell = table.cell(1, i)
        rule_cell.width = col_w
        rp = rule_cell.paragraphs[0]
        rp.paragraph_format.space_before = Pt(18)   # writing room above the rule
        rp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.right_indent = Cm(0.6)  # gap between adjacent rules
        set_border(rp, "bottom", WAIOURU_HILLS, 6, space=2)

    return table


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


marking_paragraph(section.header)
marking_paragraph(section.footer)
info = section.footer.add_paragraph()
info.paragraph_format.space_before = Pt(2)
info.paragraph_format.space_after = Pt(0)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)


def footer_run(text):
    run = info.add_run(text)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(8.5)
    force_color(run, DARKEST_HOUR)
    return run


footer_run(FOOTER_LEFT)
footer_run("\t")
footer_run(FOOTER_REFERENCE)
footer_run("\tPage ")
for r in add_field(info, "PAGE", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)
footer_run(" of ")
for r in add_field(info, "NUMPAGES", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)

# -------------------------------------------------------------- letterhead --

_logo = Image.open(LOGO_FILE).convert("RGBA")
_logo = _logo.crop(_logo.getchannel("A").getbbox())
_buf = io.BytesIO()
_logo.save(_buf, "PNG")
_buf.seek(0)
doc.add_picture(_buf, width=Mm(42))
logo_para = doc.paragraphs[-1]
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after = Pt(14)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(1)
t = title_p.add_run(TITLE)
force_font(t, FONT_DISPLAY)
t.font.size = Pt(20)
t.bold = True
force_color(t, DARKEST_HOUR)

sub_title_p = doc.add_paragraph()
sub_title_p.paragraph_format.space_after = Pt(4)
set_border(sub_title_p, "bottom", ARMY_RED, 18, space=8)
st = sub_title_p.add_run(SUBTITLE_LINE)
force_font(st, FONT_HEAD)
st.font.size = Pt(12)
st.bold = True
force_color(st, SWAMP_GREEN)

org_p = doc.add_paragraph()
org_p.paragraph_format.space_before = Pt(6)
org_p.paragraph_format.space_after = Pt(6)
o = org_p.add_run(ORIGINATOR)
force_font(o, FONT_HEAD)
o.font.size = Pt(10.5)
force_color(o, SWAMP_GREEN)

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_after = Pt(10)
set_border(meta_p, "bottom", WAIOURU_HILLS, 6, space=6)
lab = meta_p.add_run("Date  ")
force_font(lab, FONT_HEAD)
lab.font.size = Pt(9)
lab.bold = True
force_color(lab, SWAMP_GREEN)
val = meta_p.add_run(DATE)
force_font(val, FONT_HEAD)
val.font.size = Pt(9)
force_color(val, DARKEST_HOUR)

# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

action_no = 0     # counted only to report the total on build
form_labels = []  # consecutive '_' lines become one side-by-side sign-off strip


def flush_form_labels():
    global form_labels
    if form_labels:
        add_signoff_strip(form_labels)
        form_labels = []


for line in lines:
    line = line.strip()

    if line.startswith("_ "):
        form_labels.append(line[2:].strip())
        continue

    flush_form_labels()

    if not line or line.startswith("# "):
        continue          # document title already carried by the letterhead

    if line.startswith("## "):
        add_section_heading(line[3:].strip())
        continue

    if line.startswith("* "):
        action_no += 1
        add_action(line[2:].strip())
        continue

    add_body(line)

flush_form_labels()

# ----------------------------------------------------------------- finish ---

USED_STYLE_IDS = {"Normal", "Heading1", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

def _is_spacer(p):
    """Empty and carrying nothing meaningful — no image, no rule of its own."""
    if p.text.strip():
        return False
    if p._p.findall(qn("w:r") + "/" + qn("w:drawing")):
        return False
    pPr = p._p.find(qn("w:pPr"))
    return pPr is None or pPr.find(qn("w:pBdr")) is None


while doc.paragraphs and _is_spacer(doc.paragraphs[-1]):
    _el = doc.paragraphs[-1]._element
    _el.getparent().remove(_el)

# Word expects a paragraph after a trailing table; add one once stripping is done.
if doc.element.body.findall(qn("w:tbl")) and \
        doc.element.body[-2].tag == qn("w:tbl"):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

mark_update_fields(doc)

props = doc.core_properties
props.title = f"{TITLE} — {SUBTITLE_LINE}"
props.author = ORIGINATOR

doc.save(OUTPUT_DOCX)
print(f"Saved {OUTPUT_DOCX} — {action_no} actions")
