#!/usr/bin/env python3
"""
Build the 'Individual Training – Key Takeaways' one-pager as a branded
NZ Army .docx. Letterhead-style single page (no cover, no TOC); body typeset
from SOURCE_FILE on the brand system used by build_way_forward.py.
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./individual-training-key-takeaways.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/individual-training-key-takeaways.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"
DATE               = "September 2026"
ORIGINATOR         = "Individual Training Review"
VERSION            = "Working draft v0.7"

TITLE         = "Individual Training – Key Takeaways"
SUBTITLE_LINE = "Adult Learning Principles in Defence Training, Sessions 1–4"
FOOTER_LEFT   = "Individual Training – Key Takeaways"
FOOTER_REF    = "Working synthesis"

# NZ Army palette, as published in the Visual Identity Guidelines v1.0, p58.
ARMY_RED      = "D31145"   # Pantone 200 C
DARKEST_HOUR  = "000000"   # Process Black C
RUAPEHU_WHITE = "FFFFFF"
SWAMP_GREEN   = "00261B"   # Pantone 5605 C
KAWAKAWA_LEAF = "444D06"   # Pantone 5747 C
WAIOURU_HILLS = "B3A650"   # Pantone 5853 C
MOAWHANGO     = "DFD8AD"   # Pantone 5855 C

# Typography per the guidelines p59 (PC substitutes for Neue Haas Grotesk).
FONT_DISPLAY = "Arial Black"
FONT_HEAD    = "Calibri"
FONT_BODY    = "Calibri"

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _pPr(p).append(shd)


def set_border(p, side, color, sz, space=4):
    """Paragraph border. sz is in eighths of a point."""
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    idx = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > idx:
            _run(p, text[idx:m.start()], base_font, size, bold, italic, color)
        _run(p, m.group(1), base_font, size, True, italic, color)
        idx = m.end()
    if idx < len(text):
        _run(p, text[idx:], base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.header_distance = Cm(0.9)
section.footer_distance = Cm(0.9)

TEXT_WIDTH_CM = 16.6
COLUMN_WIDTH_CM = 8.3

# Six takeaways sit comfortably in one column at a readable size. Set to 2
# when the list grows enough to spill a second page.
TAKEAWAY_COLUMNS = 1

normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(9.6)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.06
nf.space_after = Pt(4)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_DISPLAY)
h1.font.bold = False
h1.font.size = Pt(12)
force_color(h1, DARKEST_HOUR)
h1.paragraph_format.space_before = Pt(10)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True


def add_section_heading(text):
    """Sections hang from the datum: a black rule above, heading beneath."""
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "top", DARKEST_HOUR, 12, space=6)
    add_text_runs(p, text, base_font=FONT_DISPLAY, size=Pt(12), bold=False,
                  color=DARKEST_HOUR)
    return p


def add_body(text, size=None):
    p = doc.add_paragraph()
    add_text_runs(p, text, size=size)
    return p


def add_callout(lines, bold=True, italic=False):
    """Shaded callout. Several lines share one block so the proposition and
    its closing distinction read as a single statement."""
    if isinstance(lines, str):
        lines = [lines]
    p = doc.add_paragraph()
    set_shading(p, MOAWHANGO)
    set_border(p, "left", ARMY_RED, 28, space=8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i:
            br = p.add_run()
            force_font(br, FONT_BODY)
            br.font.size = Pt(9.4)
            br.add_break()
        add_text_runs(p, line, size=Pt(9.4), bold=bold, italic=italic,
                      color=SWAMP_GREEN)
    return p


def add_band(text):
    """Display band: the continuum and its calibration, set apart from the
    running text so they read as the spine of the argument."""
    p = doc.add_paragraph()
    set_shading(p, SWAMP_GREEN)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.1)
    p.paragraph_format.right_indent = Cm(0.1)
    # Sized to keep the full continuum on one line: a single orphaned
    # word centred beneath the band reads as a mistake.
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(9.0), bold=True,
                  color=RUAPEHU_WHITE)
    return p


def add_takeaway(num, text, container=None):
    p = (container or doc).add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    n = p.add_run(f"{num}\t")
    force_font(n, FONT_DISPLAY)
    n.font.size = Pt(9.6)
    force_color(n, ARMY_RED)
    add_text_runs(p, text)
    return p


def _clear_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def add_takeaway_block(items):
    """Set the takeaways, in one or two columns per TAKEAWAY_COLUMNS.

    A long list spills a near-empty second page in a single column; two
    columns hold it on one page without cutting substance, with the bands and
    proposition left full width so the argument still reads down the middle.
    """
    if TAKEAWAY_COLUMNS == 1:
        for num, text in items:
            add_takeaway(num, text)
        return None

    # Split where the two columns come out closest to equal in depth. Length
    # is a good enough proxy for depth at a fixed measure.
    lengths = [len(t) for _, t in items]
    total = sum(lengths)
    best, running = 1, 0
    for i in range(1, len(items)):
        running += lengths[i - 1]
        if abs(2 * running - total) < abs(2 * sum(lengths[:best]) - total):
            best = i

    table = doc.add_table(rows=1, cols=2)
    _clear_borders(table)
    table.autofit = False
    for cell, group in zip(table.rows[0].cells,
                           (items[:best], items[best:])):
        cell.width = Cm(COLUMN_WIDTH_CM)
        cell.vertical_alignment = None
        for num, text in group:
            add_takeaway(num, text, container=cell)
        # drop the empty paragraph the cell is created with
        first = cell.paragraphs[0]
        if not first.text.strip():
            first._element.getparent().remove(first._element)
    # gutter between the columns
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for side, w in (("left", 0), ("right", 170)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)
    return table


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


marking_paragraph(section.header)
marking_paragraph(section.footer)
info = section.footer.add_paragraph()
info.paragraph_format.space_before = Pt(2)
info.paragraph_format.space_after = Pt(0)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)


def footer_run(text):
    run = info.add_run(text)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(8.5)
    force_color(run, DARKEST_HOUR)
    return run


footer_run(FOOTER_LEFT)
footer_run("\t")
footer_run(FOOTER_REF)
footer_run("\tPage ")
for r in add_field(info, "PAGE", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)
footer_run(" of ")
for r in add_field(info, "NUMPAGES", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)

# -------------------------------------------------------------- letterhead --

_logo = Image.open(LOGO_FILE).convert("RGBA")
_logo = _logo.crop(_logo.getchannel("A").getbbox())
_buf = io.BytesIO()
_logo.save(_buf, "PNG")
_buf.seek(0)
doc.add_picture(_buf, width=Mm(36))
logo_para = doc.paragraphs[-1]
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after = Pt(9)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(1)
t = title_p.add_run(TITLE)
force_font(t, FONT_DISPLAY)
t.font.size = Pt(16.5)
t.bold = False
force_color(t, DARKEST_HOUR)

sub_title_p = doc.add_paragraph()
sub_title_p.paragraph_format.space_after = Pt(4)
set_border(sub_title_p, "bottom", ARMY_RED, 18, space=8)
st = sub_title_p.add_run(SUBTITLE_LINE)
force_font(st, FONT_HEAD)
st.font.size = Pt(11)
st.bold = True
force_color(st, SWAMP_GREEN)

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_before = Pt(4)
meta_p.paragraph_format.space_after = Pt(4)
set_border(meta_p, "bottom", WAIOURU_HILLS, 6, space=6)
for i, (label, value) in enumerate([
        ("Originator", ORIGINATOR), ("Date", DATE), ("Version", VERSION)]):
    if i:
        gap = meta_p.add_run("      ")
        force_font(gap, FONT_HEAD)
    lab = meta_p.add_run(f"{label}  ")
    force_font(lab, FONT_HEAD)
    lab.font.size = Pt(9)
    lab.bold = True
    force_color(lab, SWAMP_GREEN)
    val = meta_p.add_run(value)
    force_font(val, FONT_HEAD)
    val.font.size = Pt(9)
    force_color(val, DARKEST_HOUR)

# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    if not line.strip():
        i += 1
        continue

    # Title and strapline are already carried by the letterhead.
    if line.startswith("# ") or line.startswith("**Adult Learning Principles"):
        i += 1
        continue

    if line.startswith("## "):
        add_section_heading(line[3:].strip())
        i += 1
        continue

    if line.startswith(">> "):
        add_band(line[3:].strip())
        i += 1
        continue

    if line.startswith("> "):
        block = []
        while i < len(lines) and lines[i].startswith("> "):
            block.append(lines[i][2:].strip())
            i += 1
        add_callout(block)
        continue

    m = re.match(r"^(\d+)\. (.*)$", line)
    if m:
        run = []
        while i < len(lines):
            if not lines[i].strip():
                i += 1
                continue
            m2 = re.match(r"^(\d+)\. (.*)$", lines[i])
            if not m2:
                break
            run.append((m2.group(1), m2.group(2).strip()))
            i += 1
        add_takeaway_block(run)
        continue

    # The Commander's question itself gets callout treatment.
    if line.strip() == "**What is the Army approach to individual training?**":
        add_callout(line.strip())
        i += 1
        continue

    add_body(line.strip())
    i += 1

# ----------------------------------------------------------------- finish ---

USED_STYLE_IDS = {"Normal", "Heading1", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

while doc.paragraphs and not doc.paragraphs[-1].text.strip() \
        and not doc.paragraphs[-1]._p.findall(qn("w:r") + "/" + qn("w:drawing")):
    _el = doc.paragraphs[-1]._element
    _el.getparent().remove(_el)

mark_update_fields(doc)

props = doc.core_properties
props.title = TITLE
props.author = ORIGINATOR

doc.save(OUTPUT_DOCX)
print(f"Saved {OUTPUT_DOCX}")
