#!/usr/bin/env python3
"""
Build the NZ Army 'Performance Under Pressure and Combat Mindset' capability
note as a branded .docx staff paper.

Layout/typesetting only — the source text in SOURCE_FILE is reproduced
verbatim. The single permitted editorial change: inline filename citations
(e.g. 'PUP Workbook.pdf') are lifted out of the running text and replaced
with sequential superscript markers, collected in a 'Source Documents'
list at the end.
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./way-forward.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/combat-mindset-way-forward.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"
DOCUMENT_REFERENCE = "ACS 2026"
DATE               = "August 2026"
ORIGINATOR         = "Army Command School"
VERSION            = "Draft v1.0"
DISTRIBUTION       = "COMDT ACS"
LIST_STYLE         = "bullets"

TITLE    = "Combat Mindset Framework Proposal"
SUBTITLE = "Defining, developing and assuring the human-performance capability Army requires to remain effective in combat."
FOOTER_LEFT = "Combat Mindset Framework Proposal"

# Brand palette (NZDF Visual Identity Standards — NZ Army)
ARMY_RED       = "C62026"
DARKEST_HOUR   = "000000"
RUAPEHU_WHITE  = "FFFFFF"
SWAMP_GREEN    = "002516"
KAWAKAWA_LEAF  = "3A4B00"
WAIOURU_HILLS  = "A89662"
MOAWHANGO      = "CDD2B7"

FONT_HEAD = "Arial"         # brand face: Haas Grotesk
FONT_BODY = "Book Antiqua"  # brand face: Publico (fallback Georgia)
FONT_MONO = "Consolas"

# Known source-document filenames to convert to reference markers.
CITATION_FILES = [
    "PUP Workbook.pdf",
    "LSYS (Officers) Workbook_May 26.pdf",
    "ELDA Lead Leaders Workbook (LPR).pdf",
    "Lead Self Workbook (TAD) 2026.pdf",
    "Lead Teams Workbook.pdf",
    "LDS LEAD LEADERS WORKBOOK.pdf",
]
CITATION_RE = re.compile(
    r"\s*(" + "|".join(re.escape(f) for f in CITATION_FILES) + r")"
)

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _pPr(p).append(shd)


def set_border(p, side, color, sz, space=4):
    """Paragraph border. sz is in eighths of a point."""
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    """Set a font name and strip any theme-font attributes that would
    otherwise take precedence in Word."""
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    """Insert a Word field (PAGE, NUMPAGES, TOC ...) into a paragraph."""
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ------------------------------------------------------- citation handling --

footnote_order = []  # unique filenames in order of first appearance


def citation_number(fname):
    if fname not in footnote_order:
        footnote_order.append(fname)
    return footnote_order.index(fname) + 1


def split_citations(text):
    """Return list of (kind, value) segments: ('text', str) / ('marker', n)."""
    segments = []
    pos = 0
    for m in CITATION_RE.finditer(text):
        if m.start() > pos:
            segments.append(("text", text[pos:m.start()]))
        segments.append(("marker", citation_number(m.group(1))))
        pos = m.end()
    if pos < len(text):
        segments.append(("text", text[pos:]))
    return segments


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    """Add runs for text, honouring **bold** spans and citation markers."""
    for kind, value in split_citations(text):
        if kind == "marker":
            run = p.add_run(str(value))
            run.font.superscript = True
            force_font(run, base_font)
            if size:
                run.font.size = size
            continue
        # split on **bold** spans
        idx = 0
        for m in BOLD_RE.finditer(value):
            if m.start() > idx:
                _run(p, value[idx:m.start()], base_font, size, bold, italic, color)
            _run(p, m.group(1), base_font, size, True, italic, color)
            idx = m.end()
        if idx < len(value):
            _run(p, value[idx:], base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

# Page geometry — A4
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(2.4)
section.bottom_margin = Cm(2.4)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.header_distance = Cm(1.1)
section.footer_distance = Cm(1.1)
section.different_first_page_header_footer = True

TEXT_WIDTH_CM = 16.0  # 21 - 2.5 - 2.5

# Base style — nothing may fall back to Calibri
normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(11)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.15
nf.space_after = Pt(6)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Section headings (Heading 1 keeps outline level -> TOC)
h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_HEAD)
h1.font.bold = True
h1.font.size = Pt(14)
force_color(h1, SWAMP_GREEN)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.keep_with_next = True

# Sub-headings
h2 = doc.styles["Heading 2"]
strip_style_rpr(h2)
force_font(h2, FONT_HEAD)
h2.font.bold = True
h2.font.size = Pt(11)
force_color(h2, DARKEST_HOUR)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.keep_with_next = True


def add_section_heading(text):
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "left", ARMY_RED, 24, space=8)  # Army Red section tab
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(14), bold=True,
                  color=SWAMP_GREEN)
    return p


def add_sub_heading(text):
    p = doc.add_paragraph(style="Heading 2")
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(11), bold=True,
                  color=DARKEST_HOUR)
    return p


def add_body(text):
    p = doc.add_paragraph()
    add_text_runs(p, text)
    return p


def add_display(text_lines):
    """Single-line display statement(s): indented, italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(text_lines):
        if i:
            p.add_run().add_break()
        add_text_runs(p, line, italic=True)
    return p


def add_callout(text_lines):
    """Key summation callout: Moawhango fill, Army Red left border, bold."""
    p = doc.add_paragraph()
    set_shading(p, MOAWHANGO)
    set_border(p, "left", ARMY_RED, 28, space=8)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    for i, line in enumerate(text_lines):
        if i:
            p.add_run().add_break()
        font = FONT_HEAD if "→" in line else FONT_BODY
        add_text_runs(p, line, base_font=font, bold=True)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    bullet = p.add_run("•\t")
    force_font(bullet, FONT_HEAD)
    force_color(bullet, ARMY_RED)
    bullet.bold = True
    add_text_runs(p, text)
    return p


def add_numbered(num, lead, continuation):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    n = p.add_run(f"{num}.\t")
    force_font(n, FONT_HEAD)
    n.bold = False
    add_text_runs(p, lead)
    if continuation:
        c = doc.add_paragraph()
        c.paragraph_format.left_indent = Cm(0.75)
        c.paragraph_format.space_after = Pt(6)
        add_text_runs(c, continuation)
    return p


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    set_border(p, "bottom", WAIOURU_HILLS, 6, space=1)
    return p



LANDSCAPE_TABLES = False
FLATTEN_NEXT_HEADING = False


def add_table(rows):
    global LANDSCAPE_TABLES
    ncols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.autofit = False
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), WAIOURU_HILLS)
        borders.append(el)
    tblPr.append(borders)
    for ri, tr in enumerate(t.rows):
        trPr = tr._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        if ri == 0:
            hdr = OxmlElement("w:tblHeader")
            trPr.append(hdr)
    if LANDSCAPE_TABLES and rows[0][0] == "Function":
        widths = [Cm(6.0), Cm(7.2), Cm(12.5)]
    elif ncols == 2:
        widths = [Cm(4.4), Cm(11.6)]
    elif rows[0][0] == "Function":
        widths = [Cm(3.6), Cm(4.6), Cm(7.8)]
    else:
        widths = [Cm(4.4), Cm(5.2), Cm(6.4)]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = t.cell(r, c)
            cell.width = widths[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            if r == 0:
                force_font(run, FONT_HEAD)
                run.bold = True
                run.font.size = Pt(9.5)
                force_color(run, RUAPEHU_WHITE)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), SWAMP_GREEN)
                cell._tc.get_or_add_tcPr().append(shd)
            else:
                force_font(run, FONT_BODY)
                run.font.size = Pt(9.5)
                force_color(run, DARKEST_HOUR)
    tail = doc.add_paragraph()
    tail.paragraph_format.space_after = Pt(8)


def add_page_break():
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


# Cover (first page): marking top and bottom only
marking_paragraph(section.first_page_header)
marking_paragraph(section.first_page_footer)

# Content pages: marking top centre
marking_paragraph(section.header)

# Content pages footer: marking centre + info line
def build_info_footer(footer, width_cm):
    """Marking + document info line sized to the section's text width."""
    marking_paragraph(footer)
    info = footer.add_paragraph()
    info.paragraph_format.space_before = Pt(2)
    info.paragraph_format.space_after = Pt(0)
    info.paragraph_format.tab_stops.add_tab_stop(
        Cm(width_cm / 2), WD_TAB_ALIGNMENT.CENTER)
    info.paragraph_format.tab_stops.add_tab_stop(
        Cm(width_cm), WD_TAB_ALIGNMENT.RIGHT)

    def footer_run(text, bold=False):
        run = info.add_run(text)
        force_font(run, FONT_HEAD)
        run.font.size = Pt(8.5)
        run.bold = bold
        force_color(run, DARKEST_HOUR)
        return run

    footer_run(FOOTER_LEFT)
    footer_run("\t")
    footer_run(DOCUMENT_REFERENCE)
    footer_run("\tPage ")
    for r in add_field(info, "PAGE", "1"):
        force_font(r, FONT_HEAD)
        r.font.size = Pt(8.5)
    footer_run(" of ")
    for r in add_field(info, "NUMPAGES", "1"):
        force_font(r, FONT_HEAD)
        r.font.size = Pt(8.5)


build_info_footer(section.footer, TEXT_WIDTH_CM)

# ------------------------------------------------------------------ cover ---

# Trim the logo's transparent margins so its visible ink left-aligns with
# the title text below it.
import io
from PIL import Image
_logo = Image.open(LOGO_FILE).convert("RGBA")
_logo = _logo.crop(_logo.getchannel("A").getbbox())
_buf = io.BytesIO()
_logo.save(_buf, "PNG")
_buf.seek(0)
doc.add_picture(_buf, width=Mm(60))
logo_para = doc.paragraphs[-1]
logo_para.paragraph_format.space_before = Pt(44)
logo_para.paragraph_format.space_after = Pt(64)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(2)
t = title_p.add_run("Combat Mindset")
force_font(t, FONT_HEAD)
t.font.size = Pt(36)
t.bold = True
force_color(t, DARKEST_HOUR)

sub_title_p = doc.add_paragraph()
sub_title_p.paragraph_format.space_after = Pt(10)
set_border(sub_title_p, "bottom", ARMY_RED, 20, space=10)
st = sub_title_p.add_run("Framework Proposal")
force_font(st, FONT_HEAD)
st.font.size = Pt(17)
st.bold = True
force_color(st, SWAMP_GREEN)

promise_p = doc.add_paragraph()
promise_p.paragraph_format.space_before = Pt(8)
promise_p.paragraph_format.space_after = Pt(2)
pr = promise_p.add_run("Remain effective. Act decisively.")
force_font(pr, FONT_HEAD)
pr.font.size = Pt(13)
force_color(pr, DARKEST_HOUR)

promise2_p = doc.add_paragraph()
promise2_p.paragraph_format.space_after = Pt(84)
pr2 = promise2_p.add_run("Harder to kill.")
force_font(pr2, FONT_HEAD)
pr2.font.size = Pt(13)
pr2.bold = True
force_color(pr2, ARMY_RED)

meta_rows = [
    ("Reference", DOCUMENT_REFERENCE),
    ("Date", DATE),
    ("Originator", ORIGINATOR),
    ("Version", VERSION),
]
for i, (label, value) in enumerate(meta_rows):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(3.2))
    if i == 0:
        set_border(p, "top", WAIOURU_HILLS, 6, space=6)
    if i == len(meta_rows) - 1:
        set_border(p, "bottom", WAIOURU_HILLS, 6, space=6)
    lab = p.add_run(f"{label}\t")
    force_font(lab, FONT_HEAD)
    lab.font.size = Pt(10)
    lab.bold = True
    force_color(lab, SWAMP_GREEN)
    val = p.add_run(value)
    force_font(val, FONT_HEAD)
    val.font.size = Pt(10)
    force_color(val, DARKEST_HOUR)

add_page_break()

# --------------------------------------------------------------- contents ---

toc_head = doc.add_paragraph()
toc_head.paragraph_format.space_after = Pt(12)
set_border(toc_head, "left", ARMY_RED, 24, space=8)
_r = toc_head.add_run("Contents")
force_font(_r, FONT_HEAD)
_r.font.size = Pt(14)
_r.bold = True
force_color(_r, SWAMP_GREEN)

toc_p = doc.add_paragraph()
add_field(toc_p, r'TOC \o "1-1" \h \z \u',
          "Update fields to generate the table of contents.")

# ---------------------------------------------------- the model on a page ---
# The model keeps the body margins so its heading, illustration and footer
# share the same left edge as every other section.

from docx.enum.section import WD_SECTION_START, WD_ORIENTATION

ONEPAGER_EMBED = ("/tmp/claude-0/-home-user-PUP/"
                  "2d4cec0e-a52e-5368-bb54-803c6f37698d/scratchpad/"
                  "onepager-embed.png")
model_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
model_sec.different_first_page_header_footer = False
model_sec.top_margin = Cm(2.4)
model_sec.bottom_margin = Cm(2.4)
model_sec.left_margin = Cm(2.5)
model_sec.right_margin = Cm(2.5)
model_sec.footer.is_linked_to_previous = False
build_info_footer(model_sec.footer, TEXT_WIDTH_CM)

_mh = add_section_heading("The Model on a Page")
_mh.paragraph_format.space_before = Pt(0)
doc.add_picture(ONEPAGER_EMBED, width=Cm(TEXT_WIDTH_CM))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

body_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
body_sec.different_first_page_header_footer = False
body_sec.top_margin = Cm(2.4)
body_sec.bottom_margin = Cm(2.4)
body_sec.left_margin = Cm(2.5)
body_sec.right_margin = Cm(2.5)
body_sec.footer.is_linked_to_previous = False
build_info_footer(body_sec.footer, TEXT_WIDTH_CM)

# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    if not line.strip():
        i += 1
        continue

    if line.startswith("# ") or line.strip() == f"*{SUBTITLE}*":
        i += 1  # title/subtitle already on cover
        continue

    if line.strip() == "+++":
        # Annex on a landscape page so the division-of-functions table
        # fits comfortably on one page.
        land_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        land_sec.orientation = WD_ORIENTATION.LANDSCAPE
        land_sec.page_width = Mm(297)
        land_sec.page_height = Mm(210)
        land_sec.top_margin = Cm(1.5)
        land_sec.bottom_margin = Cm(1.5)
        land_sec.left_margin = Cm(2.0)
        land_sec.right_margin = Cm(2.0)
        land_sec.different_first_page_header_footer = False
        land_sec.footer.is_linked_to_previous = False
        build_info_footer(land_sec.footer, 25.7)
        LANDSCAPE_TABLES = True
        FLATTEN_NEXT_HEADING = True
        i += 1
        continue

    if line.strip().startswith("|"):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(set(c) <= set("-: ") for c in cells):
                tbl.append(cells)
            i += 1
        add_table(tbl)
        continue

    if line.strip() == "---":
        add_divider()
        i += 1
        continue

    if line.startswith("## "):
        _h = add_section_heading(line[3:].strip())
        if FLATTEN_NEXT_HEADING:
            _h.paragraph_format.space_before = Pt(0)
            FLATTEN_NEXT_HEADING = False
        i += 1
        continue

    if line.startswith("### "):
        add_sub_heading(line[4:].strip())
        i += 1
        continue

    if line.startswith(">! "):
        block = []
        while i < len(lines) and lines[i].startswith(">! "):
            block.append(lines[i][3:].strip())
            i += 1
        add_callout(block)
        continue

    if line.startswith("> "):
        block = []
        while i < len(lines) and lines[i].startswith("> "):
            block.append(lines[i][2:].strip())
            i += 1
        add_display(block)
        continue

    if line.startswith("* "):
        add_bullet(line[2:].strip())
        i += 1
        continue

    m = re.match(r"^(\d+)\. (.*)$", line)
    if m:
        num, lead = m.group(1), m.group(2).strip()
        continuation = None
        if i + 1 < len(lines) and lines[i + 1].startswith("   ") \
                and lines[i + 1].strip():
            continuation = lines[i + 1].strip()
            i += 1
        add_numbered(num, lead, continuation)
        i += 1
        continue

    add_body(line.strip())
    i += 1

# -------------------------------------------------------- source documents --

if footnote_order:
    add_divider()
    add_section_heading("Source Documents")

for n, fname in enumerate(footnote_order, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.75))
    numr = p.add_run(f"{n}.\t")
    force_font(numr, FONT_HEAD)
    numr.bold = True
    fn = p.add_run(fname)
    force_font(fn, FONT_MONO)
    fn.font.size = Pt(10)
    tail = p.add_run(" — internal NZALC workbook.")
    force_font(tail, FONT_BODY)

# ----------------------------------------------------------------- finish ---

# Prune unused template styles so no off-palette colour or non-brand font
# remains anywhere in the file, and point document defaults at the body font.
USED_STYLE_IDS = {"Normal", "Heading1", "Heading2", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

# Drop trailing empty paragraphs so a full final page doesn't push a blank
# page after it.
while doc.paragraphs and not doc.paragraphs[-1].text.strip() \
        and not doc.paragraphs[-1]._p.findall(qn("w:r") + "/" + qn("w:drawing")):
    _el = doc.paragraphs[-1]._element
    _el.getparent().remove(_el)

mark_update_fields(doc)

props = doc.core_properties
props.title = TITLE
props.author = ORIGINATOR
props.subject = SUBTITLE

doc.save(OUTPUT_DOCX)

print(f"Saved {OUTPUT_DOCX}")
print(f"Footnote markers reference {len(footnote_order)} unique source documents:")
for n, fname in enumerate(footnote_order, start=1):
    print(f"  {n}. {fname}")
