#!/usr/bin/env python3
"""
Build the 'Commander's Training Directive: Individual Training' as a branded
NZ Army .docx staff paper.

The directive body comes from SOURCE_FILE. Annex A is drawn from the shared
approach markdown (APPROACH_FILE), the same source the one-pager and the A3
poster build from, so the direction and the approach it issues cannot drift
apart.

Bracketed text is rendered in Army Red: this is a skeleton, and every gap the
staff must complete should be visible at a glance.

    python3 build_training_directive.py
    /usr/bin/python3 convert_to_pdf.py output/individual-training-directive.docx \\
        output/individual-training-directive.pdf
"""

import io
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image

# ----------------------------------------------------------------- CONFIG ---
SOURCE_FILE        = "./individual-training-directive.md"
APPROACH_FILE      = "./individual-training-approach.md"
LOGO_FILE          = "./assets/nz-army-logo.png"
OUTPUT_DOCX        = "./output/individual-training-directive.docx"
PROTECTIVE_MARKING = "UNCLASSIFIED"

TITLE       = "Commander's Training Directive"
SUBTITLE    = "Individual Training: Interim Direction to Training Schools"
FOOTER_LEFT = "Commander's Training Directive: Individual Training"
FOOTER_REF  = "Interim"

# Letterhead block. Every value here is for the issuing staff to complete.
HEADER_FIELDS = [
    ("Reference", "[ATG reference]"),
    ("Date", "[date]"),
    ("From", "Commander, Army Training Group"),
    ("To", "[Commandants, Army training schools]"),
    ("Copy to", "[distribution]"),
]

SIGNATURE = [
    ("[Name]", None),
    ("[Rank]", None),
    ("Commander, Army Training Group", None),
]

# NZ Army palette, as published in the Visual Identity Guidelines v1.0, p58.
ARMY_RED      = "D31145"   # Pantone 200 C
DARKEST_HOUR  = "000000"   # Process Black C
SWAMP_GREEN   = "00261B"   # Pantone 5605 C
KAWAKAWA_LEAF = "444D06"   # Pantone 5747 C
WAIOURU_HILLS = "B3A650"   # Pantone 5853 C
MOAWHANGO     = "DFD8AD"   # Pantone 5855 C

# Typography per the guidelines p59 (PC substitutes for Neue Haas Grotesk).
FONT_DISPLAY = "Arial Black"
FONT_HEAD    = "Calibri"
FONT_BODY    = "Calibri"

# ------------------------------------------------------------ XML helpers ---


def _pPr(p):
    return p._p.get_or_add_pPr()


def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _pPr(p).append(shd)


def set_cell_shading(cell, fill):
    """Fill the whole cell. Paragraph shading only paints behind the text."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_border(p, side, color, sz, space=4):
    """Paragraph border. sz is in eighths of a point."""
    pPr = _pPr(p)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def force_font(style_or_run, name):
    style_or_run.font.name = name
    rPr = style_or_run.font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:cs"), name)


def force_color(style_or_run, hexval):
    style_or_run.font.color.rgb = RGBColor.from_string(hexval)
    rPr = style_or_run.font.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeShade"), None)
        color.attrib.pop(qn("w:themeTint"), None)


def strip_style_rpr(style):
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        style.element.remove(rPr)


def add_field(p, code, placeholder=None):
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    r4 = p.add_run(placeholder if placeholder else "")
    r5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)
    return [r1, r2, r3, r4, r5]


def mark_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# --------------------------------------------------------------- text runs ---

# **bold** spans and [placeholders to be completed], in one pass
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\])")


def add_text_runs(p, text, base_font=FONT_BODY, size=None, bold=False,
                  italic=False, color=None):
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            _run(p, tok[2:-2], base_font, size, True, italic, color)
        elif tok.startswith("[") and tok.endswith("]"):
            # a gap for the issuing staff: flagged, not silently blank
            _run(p, tok, base_font, size, bold, italic, ARMY_RED)
        else:
            _run(p, tok, base_font, size, bold, italic, color)


def _run(p, text, font, size, bold, italic, color):
    run = p.add_run(text)
    force_font(run, font)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        force_color(run, color)
    return run


# ------------------------------------------------------------- doc set-up ---

doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(2.3)
section.right_margin = Cm(2.3)
section.header_distance = Cm(0.9)
section.footer_distance = Cm(0.9)

TEXT_WIDTH_CM = 16.4

normal = doc.styles["Normal"]
strip_style_rpr(normal)
force_font(normal, FONT_BODY)
normal.font.size = Pt(10.2)
force_color(normal, DARKEST_HOUR)
nf = normal.paragraph_format
nf.line_spacing = 1.08
nf.space_after = Pt(5)
nf.space_before = Pt(0)
nf.alignment = WD_ALIGN_PARAGRAPH.LEFT

h1 = doc.styles["Heading 1"]
strip_style_rpr(h1)
force_font(h1, FONT_DISPLAY)
h1.font.bold = False
h1.font.size = Pt(11.5)
force_color(h1, DARKEST_HOUR)
h1.paragraph_format.space_before = Pt(11)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True

h2 = doc.styles["Heading 2"]
strip_style_rpr(h2)
force_font(h2, FONT_HEAD)
h2.font.bold = True
h2.font.size = Pt(10.5)
force_color(h2, SWAMP_GREEN)
h2.paragraph_format.space_before = Pt(8)
h2.paragraph_format.space_after = Pt(3)
h2.paragraph_format.keep_with_next = True


def add_section_heading(text):
    """Sections hang from the datum: a black rule above, heading beneath."""
    p = doc.add_paragraph(style="Heading 1")
    set_border(p, "top", DARKEST_HOUR, 12, space=6)
    add_text_runs(p, text, base_font=FONT_DISPLAY, size=Pt(11.5), bold=False,
                  color=DARKEST_HOUR)
    return p


def add_sub_heading(text):
    p = doc.add_paragraph(style="Heading 2")
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(10.5), bold=True,
                  color=SWAMP_GREEN)
    return p


def add_body(text):
    p = doc.add_paragraph()
    add_text_runs(p, text)
    return p


def add_numbered(label, text, indent=Cm(0.9), hang=None):
    """Directive paragraph: hanging label, text aligned to the tab.

    `hang` sets where the label itself sits, so sub-paragraphs can hang under
    the text of the paragraph they belong to rather than at the margin.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = -(indent - (hang or Cm(0)))
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.tab_stops.add_tab_stop(indent)
    n = p.add_run(f"{label}\t")
    force_font(n, FONT_BODY)
    add_text_runs(p, text)
    return p


def add_intent(lines):
    """Intent, set apart. It is the part a Commandant must not miss."""
    p = doc.add_paragraph()
    set_shading(p, MOAWHANGO)
    set_border(p, "left", ARMY_RED, 28, space=8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        if i:
            br = p.add_run()
            force_font(br, FONT_BODY)
            br.font.size = Pt(10)
            br.add_break()
        add_text_runs(p, line, size=Pt(10), color=SWAMP_GREEN)
    return p


def add_band(text):
    """Display band, as used for the same elements on the poster."""
    p = doc.add_paragraph()
    set_shading(p, SWAMP_GREEN)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    add_text_runs(p, text, base_font=FONT_HEAD, size=Pt(9.2), bold=True,
                  color="FFFFFF")
    return p


def add_page_break():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_signature_block():
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(16)
    spacer.paragraph_format.space_after = Pt(0)
    rule = doc.add_paragraph()
    rule.paragraph_format.left_indent = Cm(9.4)
    rule.paragraph_format.space_after = Pt(3)
    set_border(rule, "top", DARKEST_HOUR, 8, space=2)
    for text, _ in SIGNATURE:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(9.4)
        p.paragraph_format.space_after = Pt(0)
        add_text_runs(p, text, size=Pt(10))


def add_table(rows):
    """Annex B return format."""
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Cm(2.8), Cm(3.4), Cm(3.9), Cm(3.1), Cm(3.2)]
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), WAIOURU_HILLS)
        borders.append(el)
    tblPr.append(borders)

    for r, row in enumerate(rows):
        if r:
            # a blank return row has no text to give it height, so ask for it
            trPr = table.rows[r]._tr.get_or_add_trPr()
            h = OxmlElement("w:trHeight")
            h.set(qn("w:val"), "900")          # twips, about 16mm
            h.set(qn("w:hRule"), "atLeast")
            trPr.append(h)
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            if c < len(widths):
                cell.width = widths[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r == 0:
                set_cell_shading(cell, SWAMP_GREEN)
                add_text_runs(p, cell_text, base_font=FONT_HEAD, size=Pt(8.6),
                              bold=True, color="FFFFFF")
            else:
                add_text_runs(p, cell_text, size=Pt(9))
    return table


# ------------------------------------------------------- headers & footers --


def marking_paragraph(container, existing=True):
    p = container.paragraphs[0] if existing and container.paragraphs else \
        container.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROTECTIVE_MARKING)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(10)
    run.bold = True
    force_color(run, DARKEST_HOUR)
    return p


marking_paragraph(section.header)
marking_paragraph(section.footer)
info = section.footer.add_paragraph()
info.paragraph_format.space_before = Pt(2)
info.paragraph_format.space_after = Pt(0)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
info.paragraph_format.tab_stops.add_tab_stop(
    Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)


def footer_run(text):
    run = info.add_run(text)
    force_font(run, FONT_HEAD)
    run.font.size = Pt(8.5)
    force_color(run, DARKEST_HOUR)
    return run


footer_run(FOOTER_LEFT)
footer_run("\t")
footer_run(FOOTER_REF)
footer_run("\tPage ")
for r in add_field(info, "PAGE", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)
footer_run(" of ")
for r in add_field(info, "NUMPAGES", "1"):
    force_font(r, FONT_HEAD)
    r.font.size = Pt(8.5)

# -------------------------------------------------------------- letterhead --

_logo = Image.open(LOGO_FILE).convert("RGBA")
_logo = _logo.crop(_logo.getchannel("A").getbbox())
_buf = io.BytesIO()
_logo.save(_buf, "PNG")
_buf.seek(0)
doc.add_picture(_buf, width=Mm(36))
logo_para = doc.paragraphs[-1]
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after = Pt(9)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(1)
t = title_p.add_run(TITLE)
force_font(t, FONT_DISPLAY)
t.font.size = Pt(17)
force_color(t, DARKEST_HOUR)

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(5)
set_border(sub_p, "bottom", ARMY_RED, 18, space=8)
st = sub_p.add_run(SUBTITLE)
force_font(st, FONT_HEAD)
st.font.size = Pt(11)
st.bold = True
force_color(st, SWAMP_GREEN)

for label, value in HEADER_FIELDS:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(2.2)
    p.paragraph_format.first_line_indent = Cm(-2.2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(2.2))
    lab = p.add_run(f"{label}\t")
    force_font(lab, FONT_HEAD)
    lab.font.size = Pt(9.4)
    lab.bold = True
    force_color(lab, SWAMP_GREEN)
    add_text_runs(p, value, base_font=FONT_HEAD, size=Pt(9.4))

rule_p = doc.add_paragraph()
rule_p.paragraph_format.space_before = Pt(4)
rule_p.paragraph_format.space_after = Pt(0)
set_border(rule_p, "bottom", WAIOURU_HILLS, 6, space=4)

# ------------------------------------------------------------ approach pull --


def load_approach(path):
    """Proposition and the three model elements, from the shared source."""
    bands, proposition = [], []
    for line in open(path, encoding="utf-8"):
        s = line.strip()
        if s.startswith(">> "):
            bands.append(s[3:].strip())
        elif s.startswith("> "):
            proposition.append(s[2:].strip())
    if len(bands) != 3:
        raise SystemExit(f"expected 3 model elements, found {len(bands)}")
    return proposition, bands


APPROACH_LEAD = [
    ("The progression.", "Courses are described against these five stages. "
     "Instructor direction decreases across them while learner autonomy, "
     "complexity and pressure increase."),
    ("The decision rule.", "Applied to every course; the common reasoning, "
     "not a common answer."),
    ("The career cycle.", "Individual training does not end with a course. "
     "Learning is accessible before, during and after formal training."),
]


def render_approach():
    proposition, bands = load_approach(APPROACH_FILE)
    add_body("The approach issued by this directive, in full, as stated in "
             "The Army Approach to Individual Training.")
    add_intent(proposition)
    for (lead, gloss), band in zip(APPROACH_LEAD, bands):
        add_sub_heading(lead)
        add_body(gloss)
        add_band(band)


# ------------------------------------------------------------------- body ---

with open(SOURCE_FILE, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i]
    s = line.strip()

    if not s:
        i += 1
        continue

    if s.startswith("# ") or s.startswith("**Individual Training:"):
        i += 1                      # already carried by the letterhead
        continue

    if s == "<<<PAGE>>>":
        add_page_break()
        i += 1
        continue

    if s == "<<<SIGNATURE>>>":
        add_signature_block()
        i += 1
        continue

    if s == "<<<APPROACH>>>":
        render_approach()
        i += 1
        continue

    if s.startswith("### "):
        add_sub_heading(s[4:].strip())
        i += 1
        continue

    if s.startswith("## "):
        add_section_heading(s[3:].strip())
        i += 1
        continue

    if s.startswith("> "):
        block = []
        while i < len(lines) and lines[i].strip().startswith("> "):
            block.append(lines[i].strip()[2:].strip())
            i += 1
        add_intent(block)
        continue

    if s.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            # a separator row is all dashes; an empty row is a blank line of
            # the form to be written on, and must survive
            separator = all(c and set(c) <= set("- ") and "-" in c
                            for c in cells)
            if not separator:
                rows.append(cells)
            i += 1
        add_table(rows)
        continue

    m = re.match(r"^(\d+)\. (.*)$", s)
    if m:
        add_numbered(f"{m.group(1)}.", m.group(2).strip())
        i += 1
        continue

    m = re.match(r"^([a-z])\. (.*)$", s)
    if m:
        add_numbered(f"{m.group(1)}.", m.group(2).strip(),
                     indent=Cm(1.8), hang=Cm(0.9))
        i += 1
        continue

    m = re.match(r"^([A-Z])\. (.*)$", s)
    if m:
        add_numbered(f"{m.group(1)}.", m.group(2).strip())
        i += 1
        continue

    add_body(s)
    i += 1

# ----------------------------------------------------------------- finish ---

USED_STYLE_IDS = {"Normal", "Heading1", "Heading2", "DefaultParagraphFont",
                  "TableNormal", "NoList", "Header", "Footer"}
styles_el = doc.styles.element
for st in list(styles_el.findall(qn("w:style"))):
    if st.get(qn("w:styleId")) not in USED_STYLE_IDS:
        styles_el.remove(st)
for rFonts in styles_el.iter(qn("w:rFonts")):
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_BODY)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme",
                 "eastAsia"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)

mark_update_fields(doc)

props = doc.core_properties
props.title = TITLE
props.author = "Army Training Group"

doc.save(OUTPUT_DOCX)
print(f"Saved {OUTPUT_DOCX}")
